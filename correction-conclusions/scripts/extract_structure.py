#!/usr/bin/env python3
"""Extract the methodological structure of a Belgian legal conclusions .docx.

Usage:
    python extract_structure.py path/to/conclusions.docx [--out structure.json]

Outputs a JSON with:
- header_excerpt: first ~30 paragraphs concatenated
- counsel_party_detected: which party Roland Huberty / Iuxta Legal represents
- sections: detected high-level sections (faits, demandes, discussion, dispositif, inventaire)
- piece_references: every "pièce n° X" found in the body, with paragraph index
- italic_quoted_citations: runs that are both italic and contain quote marks
- faits_dates: dates found in the "Les faits" section (to check chronology)
- discussion_numbering: numbered headers detected inside the discussion
- inventaire_items: parsed items from the inventory section
- has_page_numbers: whether the document footer carries a page number

Intended to be called by Claude (skill: correction-conclusions) before audit.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "python-docx manquant. Installer : pip install -r scripts/requirements.txt\n"
    )
    sys.exit(2)


SECTION_PATTERNS = {
    "faits": re.compile(
        r"^\s*(?:I+\.?|1\.?)?\s*(?:LES\s+)?FAITS\b",
        re.IGNORECASE,
    ),
    "demandes": re.compile(
        r"^\s*(?:II+\.?|2\.?)?\s*(?:OBJET\s+DES\s+DEMANDES|DEMANDES|PROC[ÉE]DURE)\b",
        re.IGNORECASE,
    ),
    "discussion": re.compile(
        r"^\s*(?:III+\.?|3\.?)?\s*DISCUSSION\b",
        re.IGNORECASE,
    ),
    "dispositif": re.compile(
        r"^\s*(?:PAR\s+CES\s+MOTIFS|DISPOSITIF)\b",
        re.IGNORECASE,
    ),
    "inventaire": re.compile(
        r"^\s*(?:INVENTAIRE(?:\s+DES\s+PI[ÈE]CES)?|BORDEREAU\s+DES\s+PI[ÈE]CES)\b",
        re.IGNORECASE,
    ),
}

PIECE_PATTERN = re.compile(
    r"pi[èe]ces?\s+(?:n[°os]\.?\s*)?(\d+(?:[\.,]\d+)?(?:\s*(?:[àa]|et|,)\s*\d+(?:[\.,]\d+)?)*)",
    re.IGNORECASE,
)

DATE_PATTERN = re.compile(
    r"\b(?P<d>\d{1,2})[\s\./-](?P<m>\d{1,2}|janvier|f[ée]vrier|mars|avril|mai|juin|"
    r"juillet|ao[ûu]t|septembre|octobre|novembre|d[ée]cembre)[\s\./-](?P<y>\d{2,4})\b",
    re.IGNORECASE,
)

NUMBERING_PATTERN = re.compile(
    r"^\s*(?P<marker>"
    r"[A-Z]\.\d+(?:\.\d+)?|"   # A.1, A.1.2
    r"[A-Z]\.|"                 # A.
    r"\d+\.\d+(?:\.\d+)?|"     # 1.1, 1.1.2
    r"\d+\.|"                   # 1.
    r"\([a-z]\)|"               # (a)
    r"[IVX]+\.|"                # I., II., III.
    r"§\s*\d+"                  # §1
    r")\s+"
)

INVENTAIRE_ITEM_PATTERN = re.compile(
    r"^\s*(?P<num>\d+(?:\.\d+)?)\s*[\.\)\-:]\s*(?P<desc>.+)$"
)

MONTHS = {
    "janvier": 1, "janv": 1,
    "février": 2, "fevrier": 2, "févr": 2, "fevr": 2,
    "mars": 3,
    "avril": 4, "avr": 4,
    "mai": 5,
    "juin": 6,
    "juillet": 7, "juill": 7, "juil": 7,
    "août": 8, "aout": 8,
    "septembre": 9, "sept": 9,
    "octobre": 10, "oct": 10,
    "novembre": 11, "nov": 11,
    "décembre": 12, "decembre": 12, "déc": 12, "dec": 12,
}


def parse_date(d_str: str, m_str: str, y_str: str) -> tuple[int, int, int] | None:
    try:
        day = int(d_str)
        if m_str.isdigit():
            month = int(m_str)
        else:
            key = m_str.lower().rstrip(".")
            month = MONTHS.get(key)
            if month is None:
                return None
        year = int(y_str)
        if year < 100:
            year += 2000 if year < 50 else 1900
        if not (1 <= day <= 31 and 1 <= month <= 12 and 1900 <= year <= 2100):
            return None
        return (year, month, day)
    except (ValueError, AttributeError):
        return None


def detect_counsel_party(paragraphs: list[str], window: int = 40) -> dict | None:
    """Find the bloc "ayant pour conseil Roland Huberty / Iuxta Legal" and identify the party."""
    head = paragraphs[:window]
    party_label_pat = re.compile(
        r"\b(demandeur|demanderesse|d[ée]fendeur|d[ée]fenderesse|"
        r"intervenant(?:e)?(?:\s+volontaire|\s+forc[ée])?|"
        r"requ[ée]rant(?:e)?|appelant(?:e)?|intim[ée](?:e)?|"
        r"partie\s+civile|cit[ée]\s+en\s+intervention)"
        r"s?\b",
        re.IGNORECASE,
    )
    for i, para in enumerate(head):
        if re.search(r"Roland\s+Huberty|Iuxta\s+Legal", para, re.IGNORECASE):
            label = None
            for k in range(max(0, i - 12), i + 2):
                if k >= len(paragraphs):
                    break
                m = party_label_pat.search(paragraphs[k])
                if m:
                    label = m.group(1).lower()
            return {
                "label": label,
                "paragraph_index": i,
                "context": "\n".join(
                    paragraphs[max(0, i - 4):min(len(paragraphs), i + 2)]
                )[:600],
            }
    return None


def detect_sections(paragraphs: list[str]) -> list[dict]:
    sections: list[dict] = []
    for i, text in enumerate(paragraphs):
        clean = text.strip()
        if not clean:
            continue
        for stype, pat in SECTION_PATTERNS.items():
            if pat.search(clean):
                if any(s["start_para"] == i for s in sections):
                    continue
                sections.append(
                    {"type": stype, "start_para": i, "title": clean[:200]}
                )
                break
    sections.sort(key=lambda s: s["start_para"])
    for j, sec in enumerate(sections):
        sec["end_para"] = (
            sections[j + 1]["start_para"] - 1
            if j + 1 < len(sections)
            else len(paragraphs) - 1
        )
    return sections


def extract_italic_quoted(doc) -> list[dict]:
    """Find runs that are italic and contain quote characters (citation textuelle)."""
    out = []
    quote_chars = set('"«»“”‘’')
    for pi, para in enumerate(doc.paragraphs):
        for ri, run in enumerate(para.runs):
            txt = run.text or ""
            if not txt.strip():
                continue
            italic = run.italic
            if italic is None and run.font and run.font.italic is not None:
                italic = run.font.italic
            if italic and any(c in quote_chars for c in txt):
                out.append({
                    "paragraph": pi,
                    "run": ri,
                    "text": txt[:240],
                })
    return out


def extract_piece_refs(paragraphs: list[str]) -> list[dict]:
    refs = []
    for i, text in enumerate(paragraphs):
        for m in PIECE_PATTERN.finditer(text):
            refs.append({
                "paragraph": i,
                "match": m.group(0),
                "numbers_raw": m.group(1),
            })
    return refs


def extract_faits_dates(paragraphs: list[str], sections: list[dict]) -> list[dict]:
    faits = next((s for s in sections if s["type"] == "faits"), None)
    if not faits:
        return []
    out = []
    for i in range(faits["start_para"], faits["end_para"] + 1):
        if i >= len(paragraphs):
            break
        for m in DATE_PATTERN.finditer(paragraphs[i]):
            parsed = parse_date(m.group("d"), m.group("m"), m.group("y"))
            out.append({
                "paragraph": i,
                "match": m.group(0),
                "parsed": parsed,
            })
    return out


def check_chronology(faits_dates: list[dict]) -> dict:
    """Return whether dates appear in chronological order."""
    parsed = [d for d in faits_dates if d["parsed"] is not None]
    if len(parsed) < 2:
        return {"checked_dates": len(parsed), "violations": []}
    violations = []
    for i in range(1, len(parsed)):
        if parsed[i]["parsed"] < parsed[i - 1]["parsed"]:
            violations.append({
                "previous": {
                    "paragraph": parsed[i - 1]["paragraph"],
                    "match": parsed[i - 1]["match"],
                    "parsed": parsed[i - 1]["parsed"],
                },
                "current": {
                    "paragraph": parsed[i]["paragraph"],
                    "match": parsed[i]["match"],
                    "parsed": parsed[i]["parsed"],
                },
            })
    return {"checked_dates": len(parsed), "violations": violations}


def extract_discussion_numbering(paragraphs: list[str], sections: list[dict]) -> list[dict]:
    disc = next((s for s in sections if s["type"] == "discussion"), None)
    if not disc:
        return []
    out = []
    for i in range(disc["start_para"], disc["end_para"] + 1):
        if i >= len(paragraphs):
            break
        text = paragraphs[i]
        m = NUMBERING_PATTERN.match(text)
        if m and len(text.strip()) < 300:
            out.append({
                "paragraph": i,
                "marker": m.group("marker"),
                "title": text.strip()[:220],
            })
    return out


def extract_inventaire(paragraphs: list[str], sections: list[dict]) -> list[dict]:
    inv = next((s for s in sections if s["type"] == "inventaire"), None)
    if not inv:
        return []
    out = []
    for i in range(inv["start_para"], inv["end_para"] + 1):
        if i >= len(paragraphs):
            break
        m = INVENTAIRE_ITEM_PATTERN.match(paragraphs[i])
        if m:
            out.append({
                "paragraph": i,
                "number": m.group("num"),
                "description": m.group("desc").strip()[:300],
            })
    return out


def detect_page_numbers(doc) -> bool:
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            try:
                xml = footer._element.xml if footer is not None else ""
            except Exception:
                xml = ""
            if "PAGE" in xml or "<w:pgNum" in xml or "<w:instrText" in xml and "PAGE" in xml:
                return True
    return False


def extract(docx_path: Path) -> dict:
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs]
    sections = detect_sections(paragraphs)
    counsel = detect_counsel_party(paragraphs)
    piece_refs = extract_piece_refs(paragraphs)
    italic_quoted = extract_italic_quoted(doc)
    faits_dates = extract_faits_dates(paragraphs, sections)
    chronology = check_chronology(faits_dates)
    discussion_numbering = extract_discussion_numbering(paragraphs, sections)
    inventaire_items = extract_inventaire(paragraphs, sections)
    has_page_numbers = detect_page_numbers(doc)

    citing_piece_numbers = set()
    for ref in piece_refs:
        for token in re.split(r"[\s,;]+|[àaet]+", ref["numbers_raw"]):
            t = token.strip().replace(",", ".")
            if t and re.match(r"^\d+(\.\d+)?$", t):
                citing_piece_numbers.add(t)
    inventaire_numbers = {i["number"] for i in inventaire_items}
    pieces_cited_not_in_inventory = sorted(citing_piece_numbers - inventaire_numbers)
    pieces_in_inventory_never_cited = sorted(inventaire_numbers - citing_piece_numbers)

    return {
        "file": str(docx_path),
        "total_paragraphs": len(paragraphs),
        "header_excerpt": "\n".join(paragraphs[:30])[:3000],
        "counsel_party_detected": counsel,
        "sections": sections,
        "piece_references": piece_refs[:300],
        "italic_quoted_citations": italic_quoted[:200],
        "faits_dates": faits_dates,
        "chronology_check": chronology,
        "discussion_numbering": discussion_numbering[:200],
        "inventaire_items": inventaire_items,
        "inventaire_consistency": {
            "pieces_cited_not_in_inventory": pieces_cited_not_in_inventory,
            "pieces_in_inventory_never_cited": pieces_in_inventory_never_cited,
        },
        "has_page_numbers": has_page_numbers,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", type=Path, help="Chemin du .docx à analyser")
    parser.add_argument(
        "--out", type=Path, default=None,
        help="Fichier JSON de sortie (sinon stdout)",
    )
    args = parser.parse_args()

    if not args.docx.exists():
        sys.stderr.write(f"Fichier introuvable : {args.docx}\n")
        return 1

    result = extract(args.docx)
    output = json.dumps(result, indent=2, ensure_ascii=False)

    if args.out:
        args.out.write_text(output, encoding="utf-8")
        sys.stderr.write(f"Structure écrite dans {args.out}\n")
    else:
        print(output)
    return 0


if __name__ == "__main__":
    sys.exit(main())
