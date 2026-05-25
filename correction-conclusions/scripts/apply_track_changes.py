#!/usr/bin/env python3
"""Apply Word track changes (revisions) to a Belgian legal conclusions .docx.

Usage:
    python apply_track_changes.py input.docx operations.json output.docx

`operations.json` schema:
{
  "author": "Docuralis — Correction conclusions",
  "operations": [
    {"type": "replace", "find": "texte exact à modifier",
     "replace": "nouveau texte", "occurrence": 1},
    {"type": "delete", "find": "texte à supprimer"},
    {"type": "insert_after", "find": "texte ancre",
     "insert": "texte à insérer juste après"},
    {"type": "insert_paragraph_after", "find_in_paragraph": "texte du paragraphe ancre",
     "paragraph": "Nouveau paragraphe complet à ajouter en track-change."}
  ]
}

`occurrence` is 1-based and optional (default 1) — selects which occurrence of `find`
across paragraphs to target when the same text appears multiple times.

Each operation is rendered as a Word revision (`<w:ins>` / `<w:del>`) signed by the
declared author. The lawyer reviews them in Word and accepts/rejects one by one.

Output: a JSON report on stdout summarising what was applied or skipped.
"""
from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    sys.stderr.write(
        "python-docx manquant. Installer : pip install -r scripts/requirements.txt\n"
    )
    sys.exit(2)


def _new_id(counter: list[int]) -> str:
    counter[0] += 1
    return str(counter[0])


def _set_attr(el, ns_attr: str, value: str) -> None:
    el.set(qn(ns_attr), value)


def _make_text(text: str) -> "OxmlElement":
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    return t


def _make_run(text: str, rpr_source=None) -> "OxmlElement":
    r = OxmlElement("w:r")
    if rpr_source is not None:
        rpr = rpr_source.find(qn("w:rPr"))
        if rpr is not None:
            r.append(copy.deepcopy(rpr))
    r.append(_make_text(text))
    return r


def _wrap_ins(content_run, author: str, date: str, change_id: str) -> "OxmlElement":
    ins = OxmlElement("w:ins")
    _set_attr(ins, "w:id", change_id)
    _set_attr(ins, "w:author", author)
    _set_attr(ins, "w:date", date)
    ins.append(content_run)
    return ins


def _convert_run_to_delete(run_el, author: str, date: str, change_id: str) -> "OxmlElement":
    """Wrap an existing <w:r> in <w:del>, converting <w:t> to <w:delText>."""
    for t in list(run_el.findall(qn("w:t"))):
        del_t = OxmlElement("w:delText")
        del_t.text = t.text
        del_t.set(qn("xml:space"), "preserve")
        run_el.replace(t, del_t)
    del_el = OxmlElement("w:del")
    _set_attr(del_el, "w:id", change_id)
    _set_attr(del_el, "w:author", author)
    _set_attr(del_el, "w:date", date)
    parent = run_el.getparent()
    idx = list(parent).index(run_el)
    parent.remove(run_el)
    del_el.append(run_el)
    parent.insert(idx, del_el)
    return del_el


def _split_run_at(run, offset: int):
    """Split a python-docx Run at offset, returning the new (right) <w:r> XML element."""
    text = run.text or ""
    if offset <= 0 or offset >= len(text):
        return None
    left, right = text[:offset], text[offset:]
    run.text = left
    new_r = copy.deepcopy(run._element)
    for t in list(new_r.findall(qn("w:t"))):
        new_r.remove(t)
    new_r.append(_make_text(right))
    run._element.addnext(new_r)
    return new_r


def _locate(paragraph, target: str):
    """Return (start_run_idx, start_off, end_run_idx, end_off) for `target` within the paragraph,
    or None if not found. Indices refer to paragraph.runs at call time."""
    runs = paragraph.runs
    full = "".join(r.text or "" for r in runs)
    idx = full.find(target)
    if idx == -1:
        return None
    end = idx + len(target)
    pos = 0
    start_ri = start_off = end_ri = end_off = None
    for ri, r in enumerate(runs):
        rl = len(r.text or "")
        if start_ri is None and pos + rl > idx:
            start_ri = ri
            start_off = idx - pos
        if pos + rl >= end:
            end_ri = ri
            end_off = end - pos
            break
        pos += rl
    if start_ri is None or end_ri is None:
        return None
    return start_ri, start_off, end_ri, end_off


def _isolate_span(paragraph, target: str) -> tuple[int, int] | None:
    """Split runs so that `target` is exactly covered by a contiguous range of runs.
    Returns (start_idx, end_idx) inclusive in paragraph.runs after splitting, or None."""
    loc = _locate(paragraph, target)
    if loc is None:
        return None
    start_ri, start_off, end_ri, end_off = loc
    runs = paragraph.runs
    if end_off < len(runs[end_ri].text or ""):
        _split_run_at(runs[end_ri], end_off)
    if start_off > 0:
        _split_run_at(runs[start_ri], start_off)
    loc2 = _locate(paragraph, target)
    if loc2 is None:
        return None
    s_ri, s_off, e_ri, e_off = loc2
    if s_off != 0:
        return None
    return s_ri, e_ri


def _find_paragraph(doc, find: str, occurrence: int) -> int | None:
    seen = 0
    for i, para in enumerate(doc.paragraphs):
        if find and find in (para.text or ""):
            seen += 1
            if seen == occurrence:
                return i
    return None


def _op_replace(doc, op, author, date, counter) -> tuple[bool, str]:
    find = op.get("find", "")
    replace = op.get("replace", "")
    occurrence = int(op.get("occurrence", 1))
    if not find:
        return False, "replace_missing_find"
    p_idx = _find_paragraph(doc, find, occurrence)
    if p_idx is None:
        return False, "find_text_not_found"
    paragraph = doc.paragraphs[p_idx]
    span = _isolate_span(paragraph, find)
    if span is None:
        return False, "could_not_isolate_span"
    s_ri, e_ri = span
    runs = paragraph.runs
    rpr_source = runs[s_ri]._element
    last_del = None
    for ri in range(s_ri, e_ri + 1):
        if ri >= len(runs):
            break
        last_del = _convert_run_to_delete(runs[ri]._element, author, date, _new_id(counter))
    if replace:
        ins = _wrap_ins(_make_run(replace, rpr_source), author, date, _new_id(counter))
        if last_del is not None:
            last_del.addnext(ins)
    return True, "ok"


def _op_delete(doc, op, author, date, counter) -> tuple[bool, str]:
    op2 = dict(op)
    op2["replace"] = ""
    return _op_replace(doc, op2, author, date, counter)


def _op_insert_after(doc, op, author, date, counter) -> tuple[bool, str]:
    find = op.get("find", "")
    insert = op.get("insert", "")
    occurrence = int(op.get("occurrence", 1))
    if not find or not insert:
        return False, "insert_after_missing_field"
    p_idx = _find_paragraph(doc, find, occurrence)
    if p_idx is None:
        return False, "find_text_not_found"
    paragraph = doc.paragraphs[p_idx]
    span = _isolate_span(paragraph, find)
    if span is None:
        return False, "could_not_isolate_span"
    s_ri, e_ri = span
    runs = paragraph.runs
    if e_ri >= len(runs):
        return False, "index_out_of_range"
    rpr_source = runs[e_ri]._element
    ins = _wrap_ins(_make_run(insert, rpr_source), author, date, _new_id(counter))
    runs[e_ri]._element.addnext(ins)
    return True, "ok"


def _op_insert_paragraph_after(doc, op, author, date, counter) -> tuple[bool, str]:
    find = op.get("find_in_paragraph", "")
    paragraph_text = op.get("paragraph", "")
    occurrence = int(op.get("occurrence", 1))
    if not find or not paragraph_text:
        return False, "insert_paragraph_after_missing_field"
    p_idx = _find_paragraph(doc, find, occurrence)
    if p_idx is None:
        return False, "find_text_not_found"
    anchor = doc.paragraphs[p_idx]
    new_p = OxmlElement("w:p")
    ppr_src = anchor._element.find(qn("w:pPr"))
    if ppr_src is not None:
        new_p.append(copy.deepcopy(ppr_src))
    rpr_source = anchor.runs[-1]._element if anchor.runs else None
    new_p.append(_wrap_ins(_make_run(paragraph_text, rpr_source), author, date, _new_id(counter)))
    pPr_in_new = new_p.find(qn("w:pPr"))
    if pPr_in_new is None:
        pPr_in_new = OxmlElement("w:pPr")
        new_p.insert(0, pPr_in_new)
    rPr_for_para = pPr_in_new.find(qn("w:rPr"))
    if rPr_for_para is None:
        rPr_for_para = OxmlElement("w:rPr")
        pPr_in_new.append(rPr_for_para)
    ins_mark = OxmlElement("w:ins")
    _set_attr(ins_mark, "w:id", _new_id(counter))
    _set_attr(ins_mark, "w:author", author)
    _set_attr(ins_mark, "w:date", date)
    rPr_for_para.append(ins_mark)
    anchor._element.addnext(new_p)
    return True, "ok"


OPERATIONS = {
    "replace": _op_replace,
    "delete": _op_delete,
    "insert_after": _op_insert_after,
    "insert_paragraph_after": _op_insert_paragraph_after,
}


def apply_operations(input_path: Path, ops_data: dict, output_path: Path) -> dict:
    author = ops_data.get("author", "Docuralis — Correction conclusions")
    operations = ops_data.get("operations", [])
    date = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    doc = Document(str(input_path))
    counter = [0]

    applied = 0
    failed: list[dict] = []

    for i, op in enumerate(operations):
        otype = op.get("type")
        handler = OPERATIONS.get(otype)
        if handler is None:
            failed.append({"index": i, "op": op, "reason": f"unknown_type:{otype}"})
            continue
        try:
            ok, reason = handler(doc, op, author, date, counter)
        except Exception as exc:
            failed.append({"index": i, "op": op, "reason": f"exception:{exc!r}"})
            continue
        if ok:
            applied += 1
        else:
            failed.append({"index": i, "op": op, "reason": reason})

    doc.save(str(output_path))

    return {
        "input": str(input_path),
        "output": str(output_path),
        "author": author,
        "operations_total": len(operations),
        "applied": applied,
        "failed": failed,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("operations_json", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    if not args.input_docx.exists():
        sys.stderr.write(f"Fichier introuvable : {args.input_docx}\n")
        return 1
    if not args.operations_json.exists():
        sys.stderr.write(f"Fichier introuvable : {args.operations_json}\n")
        return 1

    ops_data = json.loads(args.operations_json.read_text(encoding="utf-8"))
    report = apply_operations(args.input_docx, ops_data, args.output_docx)
    print(json.dumps(report, indent=2, ensure_ascii=False))
    return 0 if not report["failed"] else 0


if __name__ == "__main__":
    sys.exit(main())
